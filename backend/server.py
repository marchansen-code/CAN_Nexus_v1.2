from fastapi import FastAPI, APIRouter, HTTPException, UploadFile, File, Depends, Response, Request
from fastapi.responses import JSONResponse
from dotenv import load_dotenv
from starlette.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
import os
import logging
from pathlib import Path
from pydantic import BaseModel, Field, ConfigDict
from typing import List, Optional, Dict, Any
import uuid
from datetime import datetime, timezone, timedelta
import pdfplumber
import asyncio
from passlib.context import CryptContext

ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

# Setup Fail2Ban-compatible logging for failed login attempts
LOG_DIR = ROOT_DIR / "logs"
LOG_DIR.mkdir(exist_ok=True)

# Configure auth logger for failed logins
auth_logger = logging.getLogger("auth_failures")
auth_logger.setLevel(logging.WARNING)
auth_handler = logging.FileHandler(LOG_DIR / "auth_failures.log")
auth_handler.setFormatter(logging.Formatter(
    '%(asctime)s %(levelname)s [AUTH] %(message)s',
    datefmt='%Y-%m-%d %H:%M:%S'
))
auth_logger.addHandler(auth_handler)

# MongoDB connection
mongo_url = os.environ['MONGO_URL']
client = AsyncIOMotorClient(mongo_url)
db = client[os.environ['DB_NAME']]

# Password hashing
pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")

# Default admin credentials
DEFAULT_ADMIN_EMAIL = "marc.hansen@canusa.de"
DEFAULT_ADMIN_PASSWORD = "CanusaNexus2024!"
DEFAULT_ADMIN_NAME = "Marc Hansen"

# Active editors tracking (in-memory, for production use Redis)
active_editors = {}

# Create the main app
app = FastAPI(title="CANUSA Knowledge Hub API")

# Create a router with the /api prefix
api_router = APIRouter(prefix="/api")

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# ==================== MODELS ====================

class User(BaseModel):
    model_config = ConfigDict(extra="ignore")
    user_id: str
    email: str
    name: str
    role: str = "viewer"
    is_blocked: bool = False
    group_ids: List[str] = []
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class Group(BaseModel):
    model_config = ConfigDict(extra="ignore")
    group_id: str = Field(default_factory=lambda: f"grp_{uuid.uuid4().hex[:12]}")
    name: str
    description: Optional[str] = None
    created_by: str
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class GroupCreate(BaseModel):
    name: str
    description: Optional[str] = None

class UserSession(BaseModel):
    model_config = ConfigDict(extra="ignore")
    session_id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    user_id: str
    session_token: str
    expires_at: datetime
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class Category(BaseModel):
    model_config = ConfigDict(extra="ignore")
    category_id: str = Field(default_factory=lambda: f"cat_{uuid.uuid4().hex[:12]}")
    name: str
    parent_id: Optional[str] = None
    description: Optional[str] = None
    order: int = 0
    created_by: str
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))
    updated_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class CategoryCreate(BaseModel):
    name: str
    parent_id: Optional[str] = None
    description: Optional[str] = None
    order: int = 0

class Article(BaseModel):
    model_config = ConfigDict(extra="ignore")
    article_id: str = Field(default_factory=lambda: f"art_{uuid.uuid4().hex[:12]}")
    title: str
    content: str
    category_ids: List[str] = []  # Multiple categories
    status: str = "draft"
    tags: List[str] = []
    source_document_id: Optional[str] = None
    review_date: Optional[datetime] = None
    expiry_date: Optional[datetime] = None  # Article expires and goes back to draft
    favorited_by: List[str] = []
    contact_person_id: Optional[str] = None
    visible_to_groups: List[str] = []  # Empty = visible to all, otherwise only to these groups
    is_important: bool = False
    important_until: Optional[datetime] = None  # Important marking expires
    created_by: str
    updated_by: str
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))
    updated_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))
    view_count: int = 0

class ArticleCreate(BaseModel):
    title: str
    content: str
    category_ids: List[str] = []
    status: str = "draft"
    tags: List[str] = []
    contact_person_id: Optional[str] = None
    visible_to_groups: List[str] = []
    expiry_date: Optional[datetime] = None
    is_important: bool = False
    important_until: Optional[datetime] = None

class ArticleUpdate(BaseModel):
    title: Optional[str] = None
    content: Optional[str] = None
    category_ids: Optional[List[str]] = None
    status: Optional[str] = None
    tags: Optional[List[str]] = None
    review_date: Optional[datetime] = None
    expiry_date: Optional[datetime] = None
    contact_person_id: Optional[str] = None
    visible_to_groups: Optional[List[str]] = None
    is_important: Optional[bool] = None
    important_until: Optional[datetime] = None

class Document(BaseModel):
    model_config = ConfigDict(extra="ignore")
    document_id: str = Field(default_factory=lambda: f"doc_{uuid.uuid4().hex[:12]}")
    filename: str
    original_language: Optional[str] = None
    target_language: str = "de"
    status: str = "pending"
    page_count: int = 0
    extracted_text: Optional[str] = None
    summary: Optional[str] = None
    structured_content: Optional[Dict[str, Any]] = None
    error_message: Optional[str] = None
    uploaded_by: str
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))
    processed_at: Optional[datetime] = None

class SearchResult(BaseModel):
    article_id: str
    title: str
    content_snippet: str
    score: float
    category_name: Optional[str] = None

# ==================== AUTH HELPERS ====================

def verify_password(plain_password: str, hashed_password: str) -> bool:
    return pwd_context.verify(plain_password, hashed_password)

def get_password_hash(password: str) -> str:
    return pwd_context.hash(password)

async def get_current_user(request: Request) -> User:
    """Get user from session token in cookie or Authorization header"""
    session_token = request.cookies.get("session_token")
    
    if not session_token:
        auth_header = request.headers.get("Authorization")
        if auth_header and auth_header.startswith("Bearer "):
            session_token = auth_header[7:]
    
    if not session_token:
        raise HTTPException(status_code=401, detail="Nicht authentifiziert")
    
    session = await db.user_sessions.find_one(
        {"session_token": session_token},
        {"_id": 0}
    )
    
    if not session:
        raise HTTPException(status_code=401, detail="Sitzung nicht gefunden")
    
    expires_at = session.get("expires_at")
    if isinstance(expires_at, str):
        expires_at = datetime.fromisoformat(expires_at)
    if expires_at.tzinfo is None:
        expires_at = expires_at.replace(tzinfo=timezone.utc)
    if expires_at < datetime.now(timezone.utc):
        raise HTTPException(status_code=401, detail="Sitzung abgelaufen")
    
    user = await db.users.find_one(
        {"user_id": session["user_id"]},
        {"_id": 0}
    )
    
    if not user:
        raise HTTPException(status_code=401, detail="Benutzer nicht gefunden")
    
    if user.get("is_blocked", False):
        raise HTTPException(status_code=403, detail="Ihr Konto wurde gesperrt. Bitte kontaktieren Sie einen Administrator.")
    
    return User(**user)

async def get_optional_user(request: Request) -> Optional[User]:
    try:
        return await get_current_user(request)
    except HTTPException:
        return None

# ==================== AUTH ENDPOINTS ====================

class LoginRequest(BaseModel):
    email: str
    password: str
    stay_logged_in: bool = False

class UserCreate(BaseModel):
    email: str
    password: str
    name: str
    role: str = "viewer"

class PasswordChange(BaseModel):
    new_password: str

@api_router.post("/auth/login")
async def login(login_data: LoginRequest, response: Response, request: Request):
    """Login with email and password"""
    # Get client IP for logging
    client_ip = request.headers.get("X-Forwarded-For", request.client.host if request.client else "unknown")
    if "," in client_ip:
        client_ip = client_ip.split(",")[0].strip()
    
    user = await db.users.find_one({"email": login_data.email.lower()}, {"_id": 0})
    
    if not user:
        # Log failed attempt - user not found
        auth_logger.warning(f"Failed login attempt from {client_ip} for unknown user: {login_data.email}")
        raise HTTPException(status_code=401, detail="Ungültige E-Mail oder Passwort")
    
    if not verify_password(login_data.password, user.get("password_hash", "")):
        # Log failed attempt - wrong password
        auth_logger.warning(f"Failed login attempt from {client_ip} for user: {login_data.email} (invalid password)")
        raise HTTPException(status_code=401, detail="Ungültige E-Mail oder Passwort")
    
    if user.get("is_blocked", False):
        # Log blocked user attempt
        auth_logger.warning(f"Blocked user login attempt from {client_ip} for user: {login_data.email}")
        raise HTTPException(status_code=403, detail="Ihr Konto wurde gesperrt")
    
    # Create session
    session_token = str(uuid.uuid4())
    days = 30 if login_data.stay_logged_in else 7
    expires_at = datetime.now(timezone.utc) + timedelta(days=days)
    
    session_doc = {
        "session_id": str(uuid.uuid4()),
        "user_id": user["user_id"],
        "session_token": session_token,
        "expires_at": expires_at.isoformat(),
        "created_at": datetime.now(timezone.utc).isoformat()
    }
    await db.user_sessions.insert_one(session_doc)
    
    # Set cookie
    response.set_cookie(
        key="session_token",
        value=session_token,
        httponly=True,
        secure=True,
        samesite="none",
        path="/",
        max_age=days * 24 * 60 * 60
    )
    
    # Log successful login (info level, different file)
    logging.info(f"Successful login from {client_ip} for user: {login_data.email}")
    
    # Return user without password hash
    return {
        "user_id": user["user_id"],
        "email": user["email"],
        "name": user["name"],
        "role": user["role"],
        "is_blocked": user.get("is_blocked", False)
    }

@api_router.get("/auth/me")
async def get_me(user: User = Depends(get_current_user)):
    """Get current authenticated user"""
    return user.model_dump()

@api_router.post("/auth/logout")
async def logout(request: Request, response: Response):
    """Logout user and clear session"""
    session_token = request.cookies.get("session_token")
    if session_token:
        await db.user_sessions.delete_one({"session_token": session_token})
    
    response.delete_cookie(key="session_token", path="/")
    return {"message": "Erfolgreich abgemeldet"}

# ==================== USER MANAGEMENT ENDPOINTS ====================

class RoleUpdate(BaseModel):
    role: str

@api_router.get("/users", response_model=List[Dict])
async def get_users(user: User = Depends(get_current_user)):
    """Get all users"""
    users = await db.users.find({}, {"_id": 0, "password_hash": 0}).sort("created_at", -1).to_list(1000)
    return users

@api_router.get("/users/{user_id}", response_model=Dict)
async def get_user(user_id: str, current_user: User = Depends(get_current_user)):
    """Get a specific user"""
    user = await db.users.find_one({"user_id": user_id}, {"_id": 0, "password_hash": 0})
    if not user:
        raise HTTPException(status_code=404, detail="Benutzer nicht gefunden")
    return user

@api_router.post("/users")
async def create_user(user_data: UserCreate, current_user: User = Depends(get_current_user)):
    """Create a new user (admin only)"""
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Nur Administratoren können Benutzer anlegen")
    
    # Check if email already exists
    existing = await db.users.find_one({"email": user_data.email.lower()})
    if existing:
        raise HTTPException(status_code=400, detail="E-Mail-Adresse bereits registriert")
    
    # Validate role
    if user_data.role not in ["admin", "editor", "viewer"]:
        raise HTTPException(status_code=400, detail="Ungültige Rolle")
    
    # Create user
    user_id = f"user_{uuid.uuid4().hex[:12]}"
    new_user = {
        "user_id": user_id,
        "email": user_data.email.lower(),
        "name": user_data.name,
        "password_hash": get_password_hash(user_data.password),
        "role": user_data.role,
        "is_blocked": False,
        "recently_viewed": [],
        "created_at": datetime.now(timezone.utc).isoformat()
    }
    await db.users.insert_one(new_user)
    
    return {
        "user_id": user_id,
        "email": user_data.email.lower(),
        "name": user_data.name,
        "role": user_data.role,
        "message": "Benutzer erfolgreich angelegt"
    }

@api_router.put("/users/{user_id}/role")
async def update_user_role(user_id: str, role_update: RoleUpdate, current_user: User = Depends(get_current_user)):
    """Update a user's role (admin only)"""
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Nur Administratoren können Rollen ändern")
    
    if role_update.role not in ["admin", "editor", "viewer"]:
        raise HTTPException(status_code=400, detail="Ungültige Rolle")
    
    if user_id == current_user.user_id:
        raise HTTPException(status_code=400, detail="Sie können Ihre eigene Rolle nicht ändern")
    
    result = await db.users.update_one(
        {"user_id": user_id},
        {"$set": {"role": role_update.role, "updated_at": datetime.now(timezone.utc).isoformat()}}
    )
    
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Benutzer nicht gefunden")
    
    return {"message": f"Rolle auf {role_update.role} geändert"}

@api_router.put("/users/{user_id}/password")
async def change_user_password(user_id: str, password_data: PasswordChange, current_user: User = Depends(get_current_user)):
    """Change a user's password (admin only)"""
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Nur Administratoren können Passwörter ändern")
    
    if len(password_data.new_password) < 6:
        raise HTTPException(status_code=400, detail="Passwort muss mindestens 6 Zeichen lang sein")
    
    result = await db.users.update_one(
        {"user_id": user_id},
        {"$set": {
            "password_hash": get_password_hash(password_data.new_password),
            "updated_at": datetime.now(timezone.utc).isoformat()
        }}
    )
    
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Benutzer nicht gefunden")
    
    # Invalidate all sessions for this user
    await db.user_sessions.delete_many({"user_id": user_id})
    
    return {"message": "Passwort erfolgreich geändert"}

@api_router.put("/users/{user_id}/block")
async def toggle_user_block(user_id: str, current_user: User = Depends(get_current_user)):
    """Block or unblock a user (admin only)"""
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Nur Administratoren können Benutzer sperren")
    
    if user_id == current_user.user_id:
        raise HTTPException(status_code=400, detail="Sie können sich nicht selbst sperren")
    
    user_doc = await db.users.find_one({"user_id": user_id})
    if not user_doc:
        raise HTTPException(status_code=404, detail="Benutzer nicht gefunden")
    
    new_status = not user_doc.get("is_blocked", False)
    
    await db.users.update_one(
        {"user_id": user_id},
        {"$set": {"is_blocked": new_status}}
    )
    
    if new_status:
        await db.user_sessions.delete_many({"user_id": user_id})
    
    return {
        "message": f"Benutzer {'gesperrt' if new_status else 'entsperrt'}",
        "is_blocked": new_status
    }

@api_router.delete("/users/{user_id}")
async def delete_user(user_id: str, current_user: User = Depends(get_current_user)):
    """Delete a user permanently (admin only)"""
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Nur Administratoren können Benutzer löschen")
    
    if user_id == current_user.user_id:
        raise HTTPException(status_code=400, detail="Sie können sich nicht selbst löschen")
    
    user_doc = await db.users.find_one({"user_id": user_id})
    if not user_doc:
        raise HTTPException(status_code=404, detail="Benutzer nicht gefunden")
    
    await db.user_sessions.delete_many({"user_id": user_id})
    await db.users.delete_one({"user_id": user_id})
    
    return {"message": "Benutzer gelöscht"}

# ==================== GROUP ENDPOINTS ====================

@api_router.get("/groups", response_model=List[Dict])
async def get_groups(user: User = Depends(get_current_user)):
    """Get all groups"""
    groups = await db.groups.find({}, {"_id": 0}).to_list(100)
    return groups

@api_router.post("/groups", response_model=Dict)
async def create_group(group: GroupCreate, user: User = Depends(get_current_user)):
    """Create a new group (admin only)"""
    if user.role != "admin":
        raise HTTPException(status_code=403, detail="Nur Administratoren können Gruppen erstellen")
    
    # Check if group name already exists
    existing = await db.groups.find_one({"name": group.name})
    if existing:
        raise HTTPException(status_code=400, detail="Eine Gruppe mit diesem Namen existiert bereits")
    
    group_doc = Group(
        name=group.name,
        description=group.description,
        created_by=user.user_id
    )
    doc = group_doc.model_dump()
    doc["created_at"] = doc["created_at"].isoformat()
    await db.groups.insert_one(doc)
    return {k: v for k, v in doc.items() if k != "_id"}

@api_router.put("/groups/{group_id}", response_model=Dict)
async def update_group(group_id: str, group: GroupCreate, user: User = Depends(get_current_user)):
    """Update a group (admin only)"""
    if user.role != "admin":
        raise HTTPException(status_code=403, detail="Nur Administratoren können Gruppen bearbeiten")
    
    result = await db.groups.update_one(
        {"group_id": group_id},
        {"$set": {
            "name": group.name,
            "description": group.description,
            "updated_at": datetime.now(timezone.utc).isoformat()
        }}
    )
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Gruppe nicht gefunden")
    
    updated = await db.groups.find_one({"group_id": group_id}, {"_id": 0})
    return updated

@api_router.delete("/groups/{group_id}")
async def delete_group(group_id: str, user: User = Depends(get_current_user)):
    """Delete a group (admin only)"""
    if user.role != "admin":
        raise HTTPException(status_code=403, detail="Nur Administratoren können Gruppen löschen")
    
    # Remove group from all users
    await db.users.update_many(
        {"group_ids": group_id},
        {"$pull": {"group_ids": group_id}}
    )
    
    # Remove group from all articles
    await db.articles.update_many(
        {"visible_to_groups": group_id},
        {"$pull": {"visible_to_groups": group_id}}
    )
    
    result = await db.groups.delete_one({"group_id": group_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Gruppe nicht gefunden")
    
    return {"message": "Gruppe gelöscht"}

class UserGroupUpdate(BaseModel):
    group_ids: List[str]

@api_router.put("/users/{user_id}/groups")
async def update_user_groups(user_id: str, data: UserGroupUpdate, current_user: User = Depends(get_current_user)):
    """Update user's group memberships (admin only)"""
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Nur Administratoren können Gruppenzugehörigkeiten ändern")
    
    # Verify all groups exist
    for gid in data.group_ids:
        exists = await db.groups.find_one({"group_id": gid})
        if not exists:
            raise HTTPException(status_code=400, detail=f"Gruppe {gid} nicht gefunden")
    
    result = await db.users.update_one(
        {"user_id": user_id},
        {"$set": {"group_ids": data.group_ids}}
    )
    
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Benutzer nicht gefunden")
    
    return {"message": "Gruppenzugehörigkeiten aktualisiert", "group_ids": data.group_ids}

@api_router.get("/groups/{group_id}/members")
async def get_group_members(group_id: str, user: User = Depends(get_current_user)):
    """Get all members of a group"""
    if user.role != "admin":
        raise HTTPException(status_code=403, detail="Nur Administratoren")
    
    members = await db.users.find(
        {"group_ids": group_id},
        {"_id": 0, "password_hash": 0}
    ).to_list(1000)
    return members

# ==================== CATEGORY ENDPOINTS ====================

@api_router.get("/categories", response_model=List[Dict])
async def get_categories(user: User = Depends(get_current_user)):
    """Get all categories as a tree structure"""
    categories = await db.categories.find({}, {"_id": 0}).to_list(1000)
    return categories

@api_router.post("/categories", response_model=Dict)
async def create_category(category: CategoryCreate, user: User = Depends(get_current_user)):
    """Create a new category"""
    cat_doc = Category(
        name=category.name,
        parent_id=category.parent_id,
        description=category.description,
        order=category.order,
        created_by=user.user_id
    )
    doc = cat_doc.model_dump()
    doc["created_at"] = doc["created_at"].isoformat()
    doc["updated_at"] = doc["updated_at"].isoformat()
    await db.categories.insert_one(doc)
    return {k: v for k, v in doc.items() if k != "_id"}

@api_router.put("/categories/{category_id}", response_model=Dict)
async def update_category(category_id: str, update: CategoryCreate, user: User = Depends(get_current_user)):
    """Update a category"""
    result = await db.categories.update_one(
        {"category_id": category_id},
        {"$set": {
            "name": update.name,
            "parent_id": update.parent_id,
            "description": update.description,
            "order": update.order,
            "updated_at": datetime.now(timezone.utc).isoformat()
        }}
    )
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Kategorie nicht gefunden")
    
    cat = await db.categories.find_one({"category_id": category_id}, {"_id": 0})
    return cat

@api_router.delete("/categories/{category_id}")
async def delete_category(category_id: str, user: User = Depends(get_current_user)):
    """Delete a category"""
    result = await db.categories.delete_one({"category_id": category_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Kategorie nicht gefunden")
    return {"message": "Kategorie gelöscht"}

# ==================== ARTICLE ENDPOINTS ====================

def can_user_see_article(article: dict, user: User) -> bool:
    """Check if user can see the article based on status, groups, etc."""
    # Admin can see everything
    if user.role == "admin":
        return True
    
    # Draft articles: only creator can see
    if article.get("status") == "draft":
        return article.get("created_by") == user.user_id
    
    # Group-restricted articles
    visible_groups = article.get("visible_to_groups", [])
    if visible_groups:
        user_groups = getattr(user, 'group_ids', []) or []
        if not any(g in user_groups for g in visible_groups):
            return False
    
    return True

@api_router.get("/articles", response_model=List[Dict])
async def get_articles(
    status: Optional[str] = None,
    category_id: Optional[str] = None,
    user: User = Depends(get_current_user)
):
    """Get all articles with optional filtering, respecting visibility rules"""
    # Get user's group memberships
    user_doc = await db.users.find_one({"user_id": user.user_id}, {"group_ids": 1})
    user_groups = user_doc.get("group_ids", []) if user_doc else []
    
    query = {"deleted_at": {"$exists": False}}  # Exclude soft-deleted articles
    if status:
        query["status"] = status
    if category_id:
        query["category_ids"] = category_id
    
    articles = await db.articles.find(query, {"_id": 0}).sort("updated_at", -1).to_list(1000)
    
    # Filter articles based on visibility
    filtered = []
    for art in articles:
        # Admin sees everything
        if user.role == "admin":
            filtered.append(art)
            continue
        
        # Draft: only creator sees it
        if art.get("status") == "draft":
            if art.get("created_by") == user.user_id:
                filtered.append(art)
            continue
        
        # Group-restricted
        visible_groups = art.get("visible_to_groups", [])
        if visible_groups:
            if any(g in user_groups for g in visible_groups):
                filtered.append(art)
        else:
            filtered.append(art)
    
    return filtered

@api_router.get("/articles/top-viewed")
async def get_top_viewed_articles(limit: int = 10, user: User = Depends(get_current_user)):
    """Get top viewed articles (published only, not deleted)"""
    articles = await db.articles.find(
        {"status": "published", "deleted_at": {"$exists": False}},
        {"_id": 0}
    ).sort("view_count", -1).limit(limit).to_list(limit)
    return articles

@api_router.get("/articles/by-category/{category_id}")
async def get_articles_by_category(category_id: str, user: User = Depends(get_current_user)):
    """Get articles in a specific category"""
    user_doc = await db.users.find_one({"user_id": user.user_id}, {"group_ids": 1})
    user_groups = user_doc.get("group_ids", []) if user_doc else []
    
    # Exclude soft-deleted articles
    articles = await db.articles.find(
        {"category_ids": category_id, "deleted_at": {"$exists": False}},
        {"_id": 0}
    ).sort("updated_at", -1).to_list(100)
    
    # Filter by visibility
    filtered = []
    for art in articles:
        if user.role == "admin":
            filtered.append(art)
            continue
        if art.get("status") == "draft" and art.get("created_by") != user.user_id:
            continue
        visible_groups = art.get("visible_to_groups", [])
        if visible_groups and not any(g in user_groups for g in visible_groups):
            continue
        filtered.append(art)
    
    return filtered

@api_router.get("/articles/{article_id}", response_model=Dict)
async def get_article(article_id: str, user: User = Depends(get_current_user)):
    """Get a single article"""
    article = await db.articles.find_one({"article_id": article_id}, {"_id": 0})
    if not article:
        raise HTTPException(status_code=404, detail="Artikel nicht gefunden")
    
    # Check visibility
    user_doc = await db.users.find_one({"user_id": user.user_id}, {"group_ids": 1})
    user_groups = user_doc.get("group_ids", []) if user_doc else []
    
    if user.role != "admin":
        if article.get("status") == "draft" and article.get("created_by") != user.user_id:
            raise HTTPException(status_code=403, detail="Zugriff verweigert")
        visible_groups = article.get("visible_to_groups", [])
        if visible_groups and not any(g in user_groups for g in visible_groups):
            raise HTTPException(status_code=403, detail="Zugriff verweigert")
    
    return article

@api_router.post("/articles", response_model=Dict)
async def create_article(article: ArticleCreate, user: User = Depends(get_current_user)):
    """Create a new article"""
    art_doc = Article(
        title=article.title,
        content=article.content,
        category_ids=article.category_ids,
        status=article.status,
        tags=article.tags,
        contact_person_id=article.contact_person_id,
        visible_to_groups=article.visible_to_groups,
        expiry_date=article.expiry_date,
        is_important=article.is_important,
        important_until=article.important_until,
        created_by=user.user_id,
        updated_by=user.user_id
    )
    doc = art_doc.model_dump()
    doc["created_at"] = doc["created_at"].isoformat()
    doc["updated_at"] = doc["updated_at"].isoformat()
    if doc.get("review_date"):
        doc["review_date"] = doc["review_date"].isoformat()
    if doc.get("expiry_date"):
        doc["expiry_date"] = doc["expiry_date"].isoformat()
    if doc.get("important_until"):
        doc["important_until"] = doc["important_until"].isoformat()
    
    await db.articles.insert_one(doc)
    return {k: v for k, v in doc.items() if k != "_id"}

@api_router.put("/articles/{article_id}", response_model=Dict)
async def update_article(article_id: str, update: ArticleUpdate, user: User = Depends(get_current_user)):
    """Update an article"""
    update_data = {k: v for k, v in update.model_dump().items() if v is not None}
    update_data["updated_by"] = user.user_id
    update_data["updated_at"] = datetime.now(timezone.utc).isoformat()
    
    # Convert dates to ISO strings
    for date_field in ["review_date", "expiry_date", "important_until"]:
        if update_data.get(date_field):
            update_data[date_field] = update_data[date_field].isoformat()
    
    result = await db.articles.update_one(
        {"article_id": article_id},
        {"$set": update_data}
    )
    
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Artikel nicht gefunden")
    
    article = await db.articles.find_one({"article_id": article_id}, {"_id": 0})
    return article

@api_router.delete("/articles/{article_id}")
async def delete_article(article_id: str, user: User = Depends(get_current_user)):
    """Soft delete an article - moves to trash for 30 days"""
    article = await db.articles.find_one({"article_id": article_id})
    if not article:
        raise HTTPException(status_code=404, detail="Artikel nicht gefunden")
    
    # Soft delete - set deleted_at timestamp
    await db.articles.update_one(
        {"article_id": article_id},
        {"$set": {
            "deleted_at": datetime.now(timezone.utc),
            "deleted_by": user.user_id
        }}
    )
    return {"message": "Artikel in Papierkorb verschoben"}

@api_router.get("/tags")
async def get_all_tags(user: User = Depends(get_current_user)):
    """Get all unique tags from articles"""
    pipeline = [
        {"$unwind": "$tags"},
        {"$group": {"_id": "$tags"}},
        {"$sort": {"_id": 1}}
    ]
    result = await db.articles.aggregate(pipeline).to_list(500)
    tags = [r["_id"] for r in result if r["_id"]]
    return {"tags": tags}

@api_router.get("/articles/search/linkable")
async def search_linkable_articles(q: str, limit: int = 10, user: User = Depends(get_current_user)):
    """Search articles for linking (@ mentions)"""
    if not q or len(q) < 2:
        return {"results": []}
    
    articles = await db.articles.find(
        {"title": {"$regex": q, "$options": "i"}},
        {"_id": 0, "article_id": 1, "title": 1, "status": 1}
    ).limit(limit).to_list(limit)
    
    return {"results": articles}

# ==================== SEARCH ====================

class SearchQuery(BaseModel):
    query: str
    top_k: int = 10
    category_id: Optional[str] = None

@api_router.post("/search")
async def search_articles(query: SearchQuery, user: User = Depends(get_current_user)):
    """Search articles with keyword matching"""
    if not query.query or len(query.query) < 2:
        return {"results": [], "query": query.query}
    
    # Build MongoDB query for text search
    search_terms = query.query.lower().split()
    
    # Search in title, content, summary, and tags
    or_conditions = []
    for term in search_terms:
        or_conditions.extend([
            {"title": {"$regex": term, "$options": "i"}},
            {"content": {"$regex": term, "$options": "i"}},
            {"summary": {"$regex": term, "$options": "i"}},
            {"tags": {"$regex": term, "$options": "i"}}
        ])
    
    mongo_query = {"$or": or_conditions}
    
    if query.category_id:
        mongo_query["category_id"] = query.category_id
    
    articles = await db.articles.find(mongo_query, {"_id": 0}).limit(query.top_k * 2).to_list(query.top_k * 2)
    
    # Score and rank results
    results = []
    for art in articles:
        title_lower = art["title"].lower()
        content_lower = art.get("content", "").lower()
        summary_lower = art.get("summary", "").lower()
        
        # Calculate relevance score
        score = 0
        for term in search_terms:
            if term in title_lower:
                score += 0.4  # Title matches are most important
            if term in summary_lower:
                score += 0.2
            if term in content_lower:
                score += 0.1
        
        if score == 0:
            continue
        
        # Find snippet
        snippet = art.get("summary", "")
        if not snippet:
            content = art.get("content", "")
            # Strip HTML tags for snippet
            import re
            clean_content = re.sub(r'<[^>]+>', '', content)
            for term in search_terms:
                idx = clean_content.lower().find(term)
                if idx >= 0:
                    start = max(0, idx - 50)
                    end = min(len(clean_content), idx + 150)
                    snippet = "..." + clean_content[start:end].strip() + "..."
                    break
            if not snippet:
                snippet = clean_content[:200] + "..." if len(clean_content) > 200 else clean_content
        
        # Get category name
        category_name = None
        if art.get("category_id"):
            cat = await db.categories.find_one({"category_id": art["category_id"]}, {"_id": 0, "name": 1})
            if cat:
                category_name = cat["name"]
        
        results.append({
            "article_id": art["article_id"],
            "title": art["title"],
            "content_snippet": snippet[:300],
            "score": min(score, 1.0),
            "category_name": category_name,
            "status": art.get("status", "draft"),
            "updated_at": art.get("updated_at")
        })
    
    # Sort by score
    results.sort(key=lambda x: x["score"], reverse=True)
    
    return {"results": results[:query.top_k], "query": query.query}

@api_router.get("/search/quick")
async def quick_search(q: str, limit: int = 5, user: User = Depends(get_current_user)):
    """Quick search for autocomplete - fast, lightweight results"""
    if not q or len(q) < 2:
        return {"results": []}
    
    # Simple regex search on title and summary only for speed
    articles = await db.articles.find(
        {"$or": [
            {"title": {"$regex": q, "$options": "i"}},
            {"summary": {"$regex": q, "$options": "i"}}
        ]},
        {"_id": 0, "article_id": 1, "title": 1, "summary": 1, "status": 1, "category_id": 1}
    ).limit(limit).to_list(limit)
    
    results = []
    for art in articles:
        # Get category name
        category_name = None
        if art.get("category_id"):
            cat = await db.categories.find_one({"category_id": art["category_id"]}, {"_id": 0, "name": 1})
            if cat:
                category_name = cat["name"]
        
        results.append({
            "article_id": art["article_id"],
            "title": art["title"],
            "summary": (art.get("summary") or "")[:100],
            "status": art.get("status", "draft"),
            "category_name": category_name
        })
    
    return {"results": results}

# ==================== DOCUMENT UPLOAD & PROCESSING ====================

@api_router.post("/documents/upload")
async def upload_document(
    file: UploadFile = File(...),
    target_language: str = "de",
    force: bool = False,
    user: User = Depends(get_current_user)
):
    """Upload a PDF document for processing"""
    if not file.filename.lower().endswith('.pdf'):
        raise HTTPException(status_code=400, detail="Nur PDF-Dateien sind erlaubt")
    
    existing_doc = await db.documents.find_one({"filename": file.filename})
    if existing_doc and not force:
        raise HTTPException(
            status_code=409, 
            detail=f"Eine Datei mit dem Namen '{file.filename}' existiert bereits."
        )
    
    if existing_doc and force:
        old_path = existing_doc.get("file_path")
        if old_path and os.path.exists(old_path):
            os.remove(old_path)
        await db.documents.delete_one({"document_id": existing_doc["document_id"]})
    
    content = await file.read()
    doc_id = f"doc_{uuid.uuid4().hex[:12]}"
    permanent_path = f"/tmp/pdfs/{doc_id}.pdf"
    
    os.makedirs("/tmp/pdfs", exist_ok=True)
    
    with open(permanent_path, "wb") as f:
        f.write(content)
    
    doc = Document(
        document_id=doc_id,
        filename=file.filename,
        target_language=target_language,
        status="pending",
        uploaded_by=user.user_id
    )
    doc_dict = doc.model_dump()
    doc_dict["created_at"] = doc_dict["created_at"].isoformat()
    doc_dict["file_path"] = permanent_path
    doc_dict["file_size"] = len(content)
    
    await db.documents.insert_one(doc_dict)
    
    asyncio.create_task(process_document(doc_id, permanent_path, target_language))
    
    return {
        "document_id": doc_id,
        "filename": doc.filename,
        "status": "pending",
        "message": "Dokument wird verarbeitet"
    }

async def process_document(document_id: str, file_path: str, target_language: str):
    """Process PDF document: extract text and tables"""
    try:
        await db.documents.update_one(
            {"document_id": document_id},
            {"$set": {"status": "processing"}}
        )
        
        extracted_text = ""
        html_content = ""
        page_count = 0
        extracted_tables = []
        
        with pdfplumber.open(file_path) as pdf:
            page_count = len(pdf.pages)
            for page_num, page in enumerate(pdf.pages, 1):
                html_content += f"<div class='pdf-page' data-page='{page_num}'>"
                
                text = page.extract_text()
                if text:
                    extracted_text += f"--- Seite {page_num} ---\n{text}\n\n"
                    paragraphs = text.split('\n\n')
                    for para in paragraphs:
                        if para.strip():
                            if len(para.strip()) < 100 and para.strip().isupper():
                                html_content += f"<h3>{para.strip()}</h3>"
                            else:
                                html_content += f"<p>{para.strip()}</p>"
                
                tables = page.extract_tables()
                for table_idx, table in enumerate(tables):
                    if table and len(table) > 0:
                        table_html = "<table class='w-full border-collapse my-4 text-sm'>"
                        for row_idx, row in enumerate(table):
                            if row:
                                table_html += "<tr>"
                                for cell in row:
                                    cell_content = cell if cell else ""
                                    if row_idx == 0:
                                        table_html += f"<th class='border border-slate-400 bg-slate-100 p-2 font-semibold text-left'>{cell_content}</th>"
                                    else:
                                        table_html += f"<td class='border border-slate-300 p-2'>{cell_content}</td>"
                                table_html += "</tr>"
                        table_html += "</table>"
                        extracted_tables.append({
                            "page": page_num,
                            "index": table_idx,
                            "html": table_html
                        })
                        html_content += table_html
                
                html_content += "</div>"
        
        if not extracted_text.strip():
            raise Exception("Kein Text konnte aus dem PDF extrahiert werden")
        
        structured_content = {
            "headlines": [],
            "bulletpoints": [],
            "tables": extracted_tables,
            "images": [],
            "html_content": html_content
        }
        
        await db.documents.update_one(
            {"document_id": document_id},
            {"$set": {
                "status": "completed",
                "page_count": page_count,
                "extracted_text": extracted_text,
                "structured_content": structured_content,
                "processed_at": datetime.now(timezone.utc).isoformat()
            }}
        )
        
        logger.info(f"Document {document_id} processed successfully")
        
    except Exception as e:
        logger.error(f"Document processing failed: {e}")
        await db.documents.update_one(
            {"document_id": document_id},
            {"$set": {
                "status": "failed",
                "error_message": str(e)
            }}
        )

@api_router.get("/documents", response_model=List[Dict])
async def get_documents(user: User = Depends(get_current_user)):
    """Get all documents (excluding soft-deleted)"""
    docs = await db.documents.find(
        {"deleted_at": {"$exists": False}}, 
        {"_id": 0, "temp_path": 0}
    ).sort("created_at", -1).to_list(100)
    return docs

@api_router.get("/documents/{document_id}", response_model=Dict)
async def get_document(document_id: str, user: User = Depends(get_current_user)):
    """Get document details"""
    doc = await db.documents.find_one(
        {"document_id": document_id, "deleted_at": {"$exists": False}}, 
        {"_id": 0, "temp_path": 0}
    )
    if not doc:
        raise HTTPException(status_code=404, detail="Dokument nicht gefunden")
    return doc

@api_router.delete("/documents/{document_id}")
async def delete_document(document_id: str, user: User = Depends(get_current_user)):
    """Soft delete a document - moves to trash for 30 days (admin only)"""
    if user.role != "admin":
        raise HTTPException(status_code=403, detail="Nur Administratoren können Dokumente löschen")
    
    doc = await db.documents.find_one({"document_id": document_id})
    if not doc:
        raise HTTPException(status_code=404, detail="Dokument nicht gefunden")
    
    # Soft delete - set deleted_at timestamp
    await db.documents.update_one(
        {"document_id": document_id},
        {"$set": {
            "deleted_at": datetime.now(timezone.utc),
            "deleted_by": user.user_id
        }}
    )
    return {"message": "Dokument in Papierkorb verschoben"}


# ==================== TRASH / PAPIERKORB ====================

@api_router.get("/trash")
async def get_trash(user: User = Depends(get_current_user)):
    """Get all soft-deleted items (admin only)"""
    if user.role != "admin":
        raise HTTPException(status_code=403, detail="Nur Administratoren können den Papierkorb sehen")
    
    # Get deleted articles
    deleted_articles = await db.articles.find(
        {"deleted_at": {"$exists": True}},
        {"_id": 0}
    ).sort("deleted_at", -1).to_list(100)
    
    # Get deleted documents
    deleted_documents = await db.documents.find(
        {"deleted_at": {"$exists": True}},
        {"_id": 0, "temp_path": 0}
    ).sort("deleted_at", -1).to_list(100)
    
    # Calculate days until permanent deletion (30 days)
    now = datetime.now(timezone.utc)
    for art in deleted_articles:
        deleted_at = art.get("deleted_at")
        if deleted_at:
            if isinstance(deleted_at, str):
                deleted_at = datetime.fromisoformat(deleted_at.replace("Z", "+00:00"))
            # Ensure timezone-aware
            if deleted_at.tzinfo is None:
                deleted_at = deleted_at.replace(tzinfo=timezone.utc)
            days_left = 30 - (now - deleted_at).days
            art["days_until_permanent_deletion"] = max(0, days_left)
    
    for doc in deleted_documents:
        deleted_at = doc.get("deleted_at")
        if deleted_at:
            if isinstance(deleted_at, str):
                deleted_at = datetime.fromisoformat(deleted_at.replace("Z", "+00:00"))
            # Ensure timezone-aware
            if deleted_at.tzinfo is None:
                deleted_at = deleted_at.replace(tzinfo=timezone.utc)
            days_left = 30 - (now - deleted_at).days
            doc["days_until_permanent_deletion"] = max(0, days_left)
    
    return {
        "articles": deleted_articles,
        "documents": deleted_documents
    }

@api_router.post("/trash/restore/article/{article_id}")
async def restore_article(article_id: str, user: User = Depends(get_current_user)):
    """Restore a soft-deleted article (admin only)"""
    if user.role != "admin":
        raise HTTPException(status_code=403, detail="Nur Administratoren können Artikel wiederherstellen")
    
    article = await db.articles.find_one({"article_id": article_id, "deleted_at": {"$exists": True}})
    if not article:
        raise HTTPException(status_code=404, detail="Artikel nicht im Papierkorb gefunden")
    
    await db.articles.update_one(
        {"article_id": article_id},
        {"$unset": {"deleted_at": "", "deleted_by": ""}}
    )
    return {"message": "Artikel wiederhergestellt"}

@api_router.post("/trash/restore/document/{document_id}")
async def restore_document(document_id: str, user: User = Depends(get_current_user)):
    """Restore a soft-deleted document (admin only)"""
    if user.role != "admin":
        raise HTTPException(status_code=403, detail="Nur Administratoren können Dokumente wiederherstellen")
    
    doc = await db.documents.find_one({"document_id": document_id, "deleted_at": {"$exists": True}})
    if not doc:
        raise HTTPException(status_code=404, detail="Dokument nicht im Papierkorb gefunden")
    
    await db.documents.update_one(
        {"document_id": document_id},
        {"$unset": {"deleted_at": "", "deleted_by": ""}}
    )
    return {"message": "Dokument wiederhergestellt"}

@api_router.delete("/trash/permanent/article/{article_id}")
async def permanently_delete_article(article_id: str, user: User = Depends(get_current_user)):
    """Permanently delete an article (admin only)"""
    if user.role != "admin":
        raise HTTPException(status_code=403, detail="Nur Administratoren können endgültig löschen")
    
    result = await db.articles.delete_one({"article_id": article_id, "deleted_at": {"$exists": True}})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Artikel nicht im Papierkorb gefunden")
    return {"message": "Artikel endgültig gelöscht"}

@api_router.delete("/trash/permanent/document/{document_id}")
async def permanently_delete_document(document_id: str, user: User = Depends(get_current_user)):
    """Permanently delete a document (admin only)"""
    if user.role != "admin":
        raise HTTPException(status_code=403, detail="Nur Administratoren können endgültig löschen")
    
    doc = await db.documents.find_one({"document_id": document_id, "deleted_at": {"$exists": True}})
    if not doc:
        raise HTTPException(status_code=404, detail="Dokument nicht im Papierkorb gefunden")
    
    # Delete the actual file
    file_path = doc.get("file_path") or doc.get("temp_path")
    if file_path and os.path.exists(file_path):
        try:
            os.remove(file_path)
        except Exception:
            pass
    
    await db.documents.delete_one({"document_id": document_id})
    return {"message": "Dokument endgültig gelöscht"}

@api_router.post("/trash/cleanup")
async def cleanup_trash(user: User = Depends(get_current_user)):
    """Auto-delete items older than 30 days (admin only)"""
    if user.role != "admin":
        raise HTTPException(status_code=403, detail="Nur Administratoren können den Papierkorb aufräumen")
    
    cutoff_date = datetime.now(timezone.utc) - timedelta(days=30)
    
    # Delete old articles
    deleted_articles = await db.articles.delete_many({
        "deleted_at": {"$exists": True, "$lt": cutoff_date}
    })
    
    # Delete old documents (and their files)
    old_docs = await db.documents.find({
        "deleted_at": {"$exists": True, "$lt": cutoff_date}
    }).to_list(100)
    
    for doc in old_docs:
        file_path = doc.get("file_path") or doc.get("temp_path")
        if file_path and os.path.exists(file_path):
            try:
                os.remove(file_path)
            except Exception:
                pass
    
    deleted_documents = await db.documents.delete_many({
        "deleted_at": {"$exists": True, "$lt": cutoff_date}
    })
    
    return {
        "message": "Papierkorb aufgeräumt",
        "deleted_articles": deleted_articles.deleted_count,
        "deleted_documents": deleted_documents.deleted_count
    }


@api_router.get("/documents/{document_id}/pdf-embed")
async def get_document_pdf_embed(document_id: str, token: str = None):
    """Get PDF file for iframe embedding"""
    from fastapi.responses import Response
    
    doc = await db.documents.find_one({"document_id": document_id})
    if not doc:
        raise HTTPException(status_code=404, detail="Dokument nicht gefunden")
    
    file_path = doc.get("file_path")
    if not file_path or not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="PDF-Datei nicht gefunden")
    
    with open(file_path, "rb") as f:
        content = f.read()
    
    return Response(
        content=content,
        media_type="application/pdf",
        headers={
            "Content-Disposition": "inline",
            "Cache-Control": "public, max-age=3600",
            "X-Frame-Options": "SAMEORIGIN"
        }
    )

# ==================== IMAGE UPLOAD ====================

@api_router.post("/images/upload")
async def upload_image(file: UploadFile = File(...), user: User = Depends(get_current_user)):
    """Upload an image for use in articles"""
    allowed_types = ["image/jpeg", "image/png", "image/gif", "image/webp"]
    if file.content_type not in allowed_types:
        raise HTTPException(status_code=400, detail="Nur Bilder sind erlaubt (JPEG, PNG, GIF, WebP)")
    
    content = await file.read()
    if len(content) > 10 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="Bild darf maximal 10MB groß sein")
    
    images_dir = "/tmp/images"
    os.makedirs(images_dir, exist_ok=True)
    
    ext = file.filename.split('.')[-1] if '.' in file.filename else 'jpg'
    image_id = f"img_{uuid.uuid4().hex[:12]}"
    filename = f"{image_id}.{ext}"
    file_path = f"{images_dir}/{filename}"
    
    with open(file_path, "wb") as f:
        f.write(content)
    
    image_doc = {
        "image_id": image_id,
        "filename": filename,
        "original_filename": file.filename,
        "content_type": file.content_type,
        "size": len(content),
        "file_path": file_path,
        "uploaded_by": user.user_id,
        "created_at": datetime.now(timezone.utc).isoformat()
    }
    await db.images.insert_one(image_doc)
    
    return {
        "image_id": image_id,
        "url": f"/api/images/{image_id}",
        "filename": file.filename
    }

@api_router.get("/images/{image_id}")
async def get_image(image_id: str):
    """Serve an uploaded image"""
    from fastapi.responses import Response
    
    image_doc = await db.images.find_one({"image_id": image_id})
    if not image_doc:
        raise HTTPException(status_code=404, detail="Bild nicht gefunden")
    
    file_path = image_doc.get("file_path")
    if not file_path or not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="Bilddatei nicht gefunden")
    
    with open(file_path, "rb") as f:
        content = f.read()
    
    return Response(
        content=content,
        media_type=image_doc.get("content_type", "image/jpeg"),
        headers={
            "Cache-Control": "public, max-age=31536000",
            "Content-Disposition": f"inline; filename=\"{image_doc.get('original_filename', 'image')}\""
        }
    )

# ==================== STATISTICS ====================

@api_router.get("/stats")
async def get_stats(user: User = Depends(get_current_user)):
    """Get dashboard statistics"""
    total_articles = await db.articles.count_documents({})
    published_articles = await db.articles.count_documents({"status": "published"})
    draft_articles = await db.articles.count_documents({"status": "draft"})
    review_articles = await db.articles.count_documents({"status": "review"})
    total_categories = await db.categories.count_documents({})
    total_documents = await db.documents.count_documents({})
    pending_documents = await db.documents.count_documents({"status": "pending"})
    
    recent_articles = await db.articles.find({}, {"_id": 0}).sort("updated_at", -1).limit(5).to_list(5)
    top_articles = await db.articles.find({}, {"_id": 0}).sort("view_count", -1).limit(5).to_list(5)
    
    favorite_articles = await db.articles.find(
        {"favorited_by": user.user_id},
        {"_id": 0}
    ).sort("updated_at", -1).limit(5).to_list(5)
    
    user_data = await db.users.find_one({"user_id": user.user_id}, {"_id": 0, "recently_viewed": 1})
    recently_viewed_ids = user_data.get("recently_viewed", [])[:15] if user_data else []
    recently_viewed = []
    if recently_viewed_ids:
        for article_id in recently_viewed_ids:
            article = await db.articles.find_one({"article_id": article_id}, {"_id": 0})
            if article:
                recently_viewed.append(article)
    
    user_articles_count = await db.articles.count_documents({"created_by": user.user_id})
    user_documents_count = await db.documents.count_documents({"uploaded_by": user.user_id})
    
    return {
        "total_articles": total_articles,
        "published_articles": published_articles,
        "draft_articles": draft_articles,
        "review_articles": review_articles,
        "total_categories": total_categories,
        "total_documents": total_documents,
        "pending_documents": pending_documents,
        "recent_articles": recent_articles,
        "top_articles": top_articles,
        "favorite_articles": favorite_articles,
        "recently_viewed": recently_viewed,
        "user_stats": {
            "articles_created": user_articles_count,
            "documents_uploaded": user_documents_count
        }
    }

# ==================== FAVORITES ====================

@api_router.post("/articles/{article_id}/favorite")
async def toggle_favorite(article_id: str, user: User = Depends(get_current_user)):
    """Toggle favorite status for an article"""
    article = await db.articles.find_one({"article_id": article_id}, {"_id": 0})
    if not article:
        raise HTTPException(status_code=404, detail="Artikel nicht gefunden")
    
    favorited_by = article.get("favorited_by", [])
    
    if user.user_id in favorited_by:
        await db.articles.update_one(
            {"article_id": article_id},
            {"$pull": {"favorited_by": user.user_id}}
        )
        return {"favorited": False, "message": "Aus Favoriten entfernt"}
    else:
        await db.articles.update_one(
            {"article_id": article_id},
            {"$addToSet": {"favorited_by": user.user_id}}
        )
        return {"favorited": True, "message": "Zu Favoriten hinzugefügt"}

@api_router.get("/favorites")
async def get_favorites(user: User = Depends(get_current_user)):
    """Get all favorite articles for current user"""
    articles = await db.articles.find(
        {"favorited_by": user.user_id},
        {"_id": 0}
    ).sort("updated_at", -1).to_list(100)
    return articles

# ==================== RECENTLY VIEWED ====================

@api_router.post("/articles/{article_id}/viewed")
async def mark_as_viewed(article_id: str, user: User = Depends(get_current_user)):
    """Mark article as viewed and update recently viewed list"""
    await db.users.update_one(
        {"user_id": user.user_id},
        {"$pull": {"recently_viewed": article_id}}
    )
    await db.users.update_one(
        {"user_id": user.user_id},
        {"$push": {"recently_viewed": {"$each": [article_id], "$position": 0, "$slice": 20}}}
    )
    await db.articles.update_one(
        {"article_id": article_id},
        {"$inc": {"view_count": 1}}
    )
    return {"message": "Als angesehen markiert"}

# ==================== PRESENCE / ACTIVE EDITORS ====================

@api_router.post("/articles/{article_id}/presence")
async def update_presence(article_id: str, user: User = Depends(get_current_user)):
    """Update editor presence for an article"""
    global active_editors
    
    if article_id not in active_editors:
        active_editors[article_id] = {}
    
    active_editors[article_id][user.user_id] = {
        "name": user.name,
        "timestamp": datetime.now(timezone.utc).isoformat()
    }
    
    cutoff = datetime.now(timezone.utc) - timedelta(seconds=30)
    for uid in list(active_editors[article_id].keys()):
        ts = datetime.fromisoformat(active_editors[article_id][uid]["timestamp"])
        if ts < cutoff:
            del active_editors[article_id][uid]
    
    others = {uid: info for uid, info in active_editors[article_id].items() if uid != user.user_id}
    
    return {"active_editors": list(others.values())}

@api_router.delete("/articles/{article_id}/presence")
async def remove_presence(article_id: str, user: User = Depends(get_current_user)):
    """Remove editor presence when leaving article"""
    global active_editors
    
    if article_id in active_editors and user.user_id in active_editors[article_id]:
        del active_editors[article_id][user.user_id]
    
    return {"message": "Präsenz entfernt"}

# ==================== WIDGET API ====================

@api_router.get("/widget/search")
async def widget_search(q: str, limit: int = 3):
    """Public widget search endpoint"""
    articles = await db.articles.find(
        {
            "status": "published",
            "$or": [
                {"title": {"$regex": q, "$options": "i"}},
                {"content": {"$regex": q, "$options": "i"}}
            ]
        },
        {"_id": 0, "article_id": 1, "title": 1, "summary": 1}
    ).limit(limit).to_list(limit)
    
    return {"results": articles, "query": q}

@api_router.get("/widget/article/{article_id}")
async def widget_get_article(article_id: str):
    """Public widget article endpoint"""
    article = await db.articles.find_one(
        {"article_id": article_id, "status": "published"},
        {"_id": 0, "article_id": 1, "title": 1, "content": 1, "summary": 1}
    )
    if not article:
        raise HTTPException(status_code=404, detail="Artikel nicht gefunden")
    return article

# ==================== BACKUP & RESTORE ====================

def json_serializer(obj):
    """Custom JSON serializer for objects not serializable by default json code"""
    if isinstance(obj, datetime):
        return obj.isoformat()
    raise TypeError(f"Object of type {type(obj)} is not JSON serializable")

@api_router.get("/backup/export")
async def export_backup(user: User = Depends(get_current_user)):
    """Export all data as JSON backup (admin only)"""
    if user.role != "admin":
        raise HTTPException(status_code=403, detail="Nur Administratoren können Backups erstellen")
    
    from fastapi.responses import StreamingResponse
    import json
    import io
    
    # Collect all data
    articles = await db.articles.find({}, {"_id": 0}).to_list(10000)
    categories = await db.categories.find({}, {"_id": 0}).to_list(1000)
    users = await db.users.find({}, {"_id": 0, "password_hash": 0}).to_list(1000)  # No passwords!
    documents_meta = await db.documents.find({}, {"_id": 0, "file_path": 0, "temp_path": 0}).to_list(1000)
    
    backup_data = {
        "version": "2.0.0",
        "created_at": datetime.now(timezone.utc).isoformat(),
        "created_by": user.email,
        "statistics": {
            "articles": len(articles),
            "categories": len(categories),
            "users": len(users),
            "documents": len(documents_meta)
        },
        "data": {
            "articles": articles,
            "categories": categories,
            "users": users,
            "documents_metadata": documents_meta
        }
    }
    
    # Create JSON file with custom serializer for datetime objects
    json_content = json.dumps(backup_data, ensure_ascii=False, indent=2, default=json_serializer)
    
    # Generate filename with timestamp
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"canusa_nexus_backup_{timestamp}.json"
    
    return StreamingResponse(
        io.BytesIO(json_content.encode('utf-8')),
        media_type="application/json",
        headers={
            "Content-Disposition": f"attachment; filename={filename}"
        }
    )

@api_router.get("/backup/documents")
async def export_documents_backup(user: User = Depends(get_current_user)):
    """Export all uploaded documents as ZIP backup (admin only)"""
    if user.role != "admin":
        raise HTTPException(status_code=403, detail="Nur Administratoren können Dokument-Backups erstellen")
    
    from fastapi.responses import StreamingResponse
    import zipfile
    import io
    import json
    
    # Get all documents with file paths
    documents = await db.documents.find({}, {"_id": 0}).to_list(1000)
    
    if not documents:
        raise HTTPException(status_code=404, detail="Keine Dokumente zum Exportieren vorhanden")
    
    # Create ZIP in memory
    zip_buffer = io.BytesIO()
    
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        # Add manifest file with metadata
        manifest = {
            "version": "2.0.0",
            "created_at": datetime.now(timezone.utc).isoformat(),
            "created_by": user.email,
            "documents": []
        }
        
        files_added = 0
        for doc in documents:
            file_path = doc.get("file_path") or doc.get("temp_path")
            
            if file_path and os.path.exists(file_path):
                # Add file to ZIP with original filename
                original_name = doc.get("filename", f"{doc['document_id']}.pdf")
                zip_path = f"documents/{original_name}"
                
                # Handle duplicate filenames
                counter = 1
                base_name = original_name.rsplit('.', 1)[0] if '.' in original_name else original_name
                ext = original_name.rsplit('.', 1)[1] if '.' in original_name else 'pdf'
                while zip_path in [info.filename for info in zip_file.filelist]:
                    zip_path = f"documents/{base_name}_{counter}.{ext}"
                    counter += 1
                
                zip_file.write(file_path, zip_path)
                files_added += 1
                
                # Add to manifest
                manifest["documents"].append({
                    "document_id": doc.get("document_id"),
                    "filename": doc.get("filename"),
                    "original_filename": original_name,
                    "zip_path": zip_path,
                    "uploaded_by": doc.get("uploaded_by"),
                    "created_at": doc.get("created_at"),
                    "file_size": doc.get("file_size")
                })
        
        manifest["total_files"] = files_added
        
        # Add manifest as JSON file
        manifest_json = json.dumps(manifest, ensure_ascii=False, indent=2, default=json_serializer)
        zip_file.writestr("manifest.json", manifest_json)
    
    if files_added == 0:
        raise HTTPException(status_code=404, detail="Keine Dokument-Dateien gefunden")
    
    zip_buffer.seek(0)
    
    # Generate filename with timestamp
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"canusa_nexus_documents_{timestamp}.zip"
    
    return StreamingResponse(
        zip_buffer,
        media_type="application/zip",
        headers={
            "Content-Disposition": f"attachment; filename={filename}"
        }
    )

@api_router.post("/backup/documents/import")
async def import_documents_backup(
    file: UploadFile = File(...),
    user: User = Depends(get_current_user)
):
    """Import documents from ZIP backup (admin only)"""
    if user.role != "admin":
        raise HTTPException(status_code=403, detail="Nur Administratoren können Dokument-Backups importieren")
    
    import zipfile
    import io
    import json
    
    if not file.filename.endswith('.zip'):
        raise HTTPException(status_code=400, detail="Bitte eine ZIP-Datei hochladen")
    
    content = await file.read()
    zip_buffer = io.BytesIO(content)
    
    results = {"imported": 0, "skipped": 0, "errors": 0}
    
    try:
        with zipfile.ZipFile(zip_buffer, 'r') as zip_file:
            # Create upload directory
            os.makedirs("/tmp/pdfs", exist_ok=True)
            
            # Extract and import documents
            for file_info in zip_file.filelist:
                if file_info.filename.startswith("documents/") and not file_info.is_dir():
                    original_name = file_info.filename.replace("documents/", "")
                    
                    # Check if document already exists
                    existing = await db.documents.find_one({"filename": original_name})
                    if existing:
                        results["skipped"] += 1
                        continue
                    
                    try:
                        # Extract file
                        file_content = zip_file.read(file_info.filename)
                        doc_id = f"doc_{uuid.uuid4().hex[:12]}"
                        file_path = f"/tmp/pdfs/{doc_id}.pdf"
                        
                        with open(file_path, "wb") as f:
                            f.write(file_content)
                        
                        # Create document record
                        doc_record = {
                            "document_id": doc_id,
                            "filename": original_name,
                            "status": "completed",
                            "file_path": file_path,
                            "file_size": len(file_content),
                            "uploaded_by": user.user_id,
                            "created_at": datetime.now(timezone.utc).isoformat(),
                            "imported_from_backup": True
                        }
                        
                        await db.documents.insert_one(doc_record)
                        results["imported"] += 1
                        
                    except Exception as e:
                        logger.error(f"Failed to import document {original_name}: {e}")
                        results["errors"] += 1
    
    except zipfile.BadZipFile:
        raise HTTPException(status_code=400, detail="Ungültige ZIP-Datei")
    
    return {
        "message": "Dokument-Import abgeschlossen",
        "results": results
    }

class BackupImportRequest(BaseModel):
    backup_data: Dict[str, Any]
    import_articles: bool = True
    import_categories: bool = True
    import_users: bool = True
    merge_mode: bool = True  # True = merge with existing, False = replace all

@api_router.post("/backup/import")
async def import_backup(request: BackupImportRequest, user: User = Depends(get_current_user)):
    """Import data from JSON backup (admin only)"""
    if user.role != "admin":
        raise HTTPException(status_code=403, detail="Nur Administratoren können Backups wiederherstellen")
    
    backup = request.backup_data
    
    # Validate backup structure
    if "data" not in backup:
        raise HTTPException(status_code=400, detail="Ungültiges Backup-Format: 'data' fehlt")
    
    results = {
        "articles": {"imported": 0, "skipped": 0, "errors": 0},
        "categories": {"imported": 0, "skipped": 0, "errors": 0},
        "users": {"imported": 0, "skipped": 0, "errors": 0}
    }
    
    data = backup["data"]
    
    # Import categories first (articles may reference them)
    if request.import_categories and "categories" in data:
        for cat in data["categories"]:
            try:
                existing = await db.categories.find_one({"category_id": cat.get("category_id")})
                if existing:
                    if not request.merge_mode:
                        await db.categories.replace_one({"category_id": cat["category_id"]}, cat)
                        results["categories"]["imported"] += 1
                    else:
                        results["categories"]["skipped"] += 1
                else:
                    await db.categories.insert_one(cat)
                    results["categories"]["imported"] += 1
            except Exception as e:
                logger.error(f"Category import error: {e}")
                results["categories"]["errors"] += 1
    
    # Import articles
    if request.import_articles and "articles" in data:
        for art in data["articles"]:
            try:
                existing = await db.articles.find_one({"article_id": art.get("article_id")})
                if existing:
                    if not request.merge_mode:
                        await db.articles.replace_one({"article_id": art["article_id"]}, art)
                        results["articles"]["imported"] += 1
                    else:
                        results["articles"]["skipped"] += 1
                else:
                    await db.articles.insert_one(art)
                    results["articles"]["imported"] += 1
            except Exception as e:
                logger.error(f"Article import error: {e}")
                results["articles"]["errors"] += 1
    
    # Import users (without passwords - they need to be set manually)
    if request.import_users and "users" in data:
        for usr in data["users"]:
            try:
                # Normalize email to lowercase for consistency
                if usr.get("email"):
                    usr["email"] = usr["email"].lower()
                existing = await db.users.find_one({"email": usr.get("email")})
                if existing:
                    results["users"]["skipped"] += 1
                else:
                    # Set a temporary password that must be changed
                    usr["password_hash"] = get_password_hash("TempPassword123!")
                    usr["needs_password_change"] = True
                    await db.users.insert_one(usr)
                    results["users"]["imported"] += 1
            except Exception as e:
                logger.error(f"User import error: {e}")
                results["users"]["errors"] += 1
    
    return {
        "message": "Backup-Import abgeschlossen",
        "results": results,
        "note": "Importierte Benutzer haben das temporäre Passwort 'TempPassword123!' und müssen es ändern."
    }

@api_router.get("/backup/preview")
async def preview_backup_info(user: User = Depends(get_current_user)):
    """Get current database statistics for backup preview (admin only)"""
    if user.role != "admin":
        raise HTTPException(status_code=403, detail="Nur Administratoren")
    
    return {
        "articles": await db.articles.count_documents({}),
        "categories": await db.categories.count_documents({}),
        "users": await db.users.count_documents({}),
        "documents": await db.documents.count_documents({})
    }

# ==================== ARTICLE EXPORT (PDF & DOCX) ====================

def strip_html(html_content: str) -> str:
    """Remove HTML tags and convert to plain text"""
    import re
    # Remove script and style elements
    text = re.sub(r'<script[^>]*>.*?</script>', '', html_content, flags=re.DOTALL)
    text = re.sub(r'<style[^>]*>.*?</style>', '', text, flags=re.DOTALL)
    # Convert some tags to text equivalents
    text = re.sub(r'<br\s*/?>', '\n', text)
    text = re.sub(r'</p>', '\n\n', text)
    text = re.sub(r'</h[1-6]>', '\n\n', text)
    text = re.sub(r'<li>', '• ', text)
    text = re.sub(r'</li>', '\n', text)
    # Remove all remaining tags
    text = re.sub(r'<[^>]+>', '', text)
    # Clean up whitespace
    text = re.sub(r'\n\s*\n', '\n\n', text)
    text = text.strip()
    return text

@api_router.get("/articles/{article_id}/export/pdf")
async def export_article_pdf(article_id: str, user: User = Depends(get_current_user)):
    """Export article as PDF"""
    from fastapi.responses import StreamingResponse
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    import io
    
    article = await db.articles.find_one({"article_id": article_id}, {"_id": 0})
    if not article:
        raise HTTPException(status_code=404, detail="Artikel nicht gefunden")
    
    # Get category name
    category_name = "Keine Kategorie"
    if article.get("category_id"):
        cat = await db.categories.find_one({"category_id": article["category_id"]}, {"_id": 0, "name": 1})
        if cat:
            category_name = cat["name"]
    
    # Get author name
    author_name = "Unbekannt"
    if article.get("created_by"):
        author = await db.users.find_one({"user_id": article["created_by"]}, {"_id": 0, "name": 1})
        if author:
            author_name = author["name"]
    
    # Create PDF
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=2*cm,
        leftMargin=2*cm,
        topMargin=2*cm,
        bottomMargin=2*cm
    )
    
    styles = getSampleStyleSheet()
    
    # Custom styles
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=18,
        spaceAfter=12,
        textColor='#1e293b'
    )
    
    meta_style = ParagraphStyle(
        'Meta',
        parent=styles['Normal'],
        fontSize=10,
        textColor='#64748b',
        spaceAfter=20
    )
    
    body_style = ParagraphStyle(
        'CustomBody',
        parent=styles['Normal'],
        fontSize=11,
        leading=16,
        alignment=TA_JUSTIFY,
        spaceAfter=12
    )
    
    # Build content
    story = []
    
    # Title
    story.append(Paragraph(article["title"], title_style))
    
    # Meta info
    created_at = article.get("created_at", "")
    if isinstance(created_at, str) and created_at:
        try:
            dt = datetime.fromisoformat(created_at.replace('Z', '+00:00'))
            created_at = dt.strftime("%d.%m.%Y")
        except Exception:
            pass
    
    meta_text = f"Kategorie: {category_name} | Autor: {author_name} | Erstellt: {created_at}"
    story.append(Paragraph(meta_text, meta_style))
    
    # Summary
    if article.get("summary"):
        story.append(Paragraph(f"<b>Zusammenfassung:</b> {article['summary']}", body_style))
        story.append(Spacer(1, 12))
    
    # Content
    content = strip_html(article.get("content", ""))
    paragraphs = content.split('\n\n')
    for para in paragraphs:
        if para.strip():
            # Escape special characters for ReportLab
            para = para.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            story.append(Paragraph(para, body_style))
    
    # Footer
    story.append(Spacer(1, 30))
    footer_style = ParagraphStyle('Footer', parent=styles['Normal'], fontSize=8, textColor='#94a3b8', alignment=TA_CENTER)
    story.append(Paragraph(f"Exportiert aus CANUSA Nexus am {datetime.now().strftime('%d.%m.%Y %H:%M')}", footer_style))
    
    doc.build(story)
    buffer.seek(0)
    
    # Sanitize filename
    safe_title = "".join(c for c in article["title"] if c.isalnum() or c in (' ', '-', '_')).strip()[:50]
    filename = f"{safe_title}.pdf"
    
    return StreamingResponse(
        buffer,
        media_type="application/pdf",
        headers={
            "Content-Disposition": f"attachment; filename=\"{filename}\""
        }
    )

@api_router.get("/articles/{article_id}/export/docx")
async def export_article_docx(article_id: str, user: User = Depends(get_current_user)):
    """Export article as Word document"""
    from fastapi.responses import StreamingResponse
    from docx import Document as DocxDocument
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import io
    
    article = await db.articles.find_one({"article_id": article_id}, {"_id": 0})
    if not article:
        raise HTTPException(status_code=404, detail="Artikel nicht gefunden")
    
    # Get category name
    category_name = "Keine Kategorie"
    if article.get("category_id"):
        cat = await db.categories.find_one({"category_id": article["category_id"]}, {"_id": 0, "name": 1})
        if cat:
            category_name = cat["name"]
    
    # Get author name
    author_name = "Unbekannt"
    if article.get("created_by"):
        author = await db.users.find_one({"user_id": article["created_by"]}, {"_id": 0, "name": 1})
        if author:
            author_name = author["name"]
    
    # Create Word document
    doc = DocxDocument()
    
    # Title
    title = doc.add_heading(article["title"], level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Meta info
    created_at = article.get("created_at", "")
    if isinstance(created_at, str) and created_at:
        try:
            dt = datetime.fromisoformat(created_at.replace('Z', '+00:00'))
            created_at = dt.strftime("%d.%m.%Y")
        except Exception:
            pass
    
    meta = doc.add_paragraph()
    meta_run = meta.add_run(f"Kategorie: {category_name} | Autor: {author_name} | Erstellt: {created_at}")
    meta_run.font.size = Pt(10)
    meta_run.font.color.rgb = RGBColor(100, 116, 139)
    
    doc.add_paragraph()  # Spacer
    
    # Summary
    if article.get("summary"):
        summary_para = doc.add_paragraph()
        summary_run = summary_para.add_run("Zusammenfassung: ")
        summary_run.bold = True
        summary_para.add_run(article["summary"])
        doc.add_paragraph()
    
    # Content
    content = strip_html(article.get("content", ""))
    paragraphs = content.split('\n\n')
    for para in paragraphs:
        if para.strip():
            p = doc.add_paragraph(para.strip())
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer_run = footer.add_run(f"Exportiert aus CANUSA Nexus am {datetime.now().strftime('%d.%m.%Y %H:%M')}")
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(148, 163, 184)
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    # Sanitize filename
    safe_title = "".join(c for c in article["title"] if c.isalnum() or c in (' ', '-', '_')).strip()[:50]
    filename = f"{safe_title}.docx"
    
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f"attachment; filename=\"{filename}\""
        }
    )

# ==================== ROOT ====================

@api_router.get("/")
async def root():
    return {"message": "CANUSA Knowledge Hub API", "version": "2.0.0"}

# Include the router in the main app
app.include_router(api_router)

app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.on_event("startup")
async def startup():
    """Initialize database and create default admin"""
    # Create indexes
    await db.articles.create_index([("title", "text"), ("content", "text")])
    await db.articles.create_index("status")
    await db.articles.create_index("category_ids")
    await db.users.create_index("email", unique=True)
    await db.user_sessions.create_index("session_token")
    await db.groups.create_index("name", unique=True)
    
    # Check for existing admin user
    admin_exists = await db.users.find_one({"email": DEFAULT_ADMIN_EMAIL})
    
    if admin_exists:
        # Check if user has password_hash (migration from Google Auth)
        if not admin_exists.get("password_hash"):
            logger.info(f"Migrating admin user {DEFAULT_ADMIN_EMAIL} to password auth...")
            await db.users.update_one(
                {"email": DEFAULT_ADMIN_EMAIL},
                {"$set": {
                    "password_hash": get_password_hash(DEFAULT_ADMIN_PASSWORD),
                    "role": "admin",
                    "is_blocked": False,
                    "group_ids": []
                }}
            )
            logger.info("Admin user migrated successfully")
    else:
        # Create new admin user
        admin_user = {
            "user_id": f"user_{uuid.uuid4().hex[:12]}",
            "email": DEFAULT_ADMIN_EMAIL,
            "name": DEFAULT_ADMIN_NAME,
            "password_hash": get_password_hash(DEFAULT_ADMIN_PASSWORD),
            "role": "admin",
            "is_blocked": False,
            "group_ids": [],
            "recently_viewed": [],
            "created_at": datetime.now(timezone.utc).isoformat()
        }
        await db.users.insert_one(admin_user)
        logger.info(f"Default admin user created: {DEFAULT_ADMIN_EMAIL}")
    
    # Process expired articles and important markings
    await process_expirations()

async def process_expirations():
    """Check for expired articles and important markings"""
    now = datetime.now(timezone.utc).isoformat()
    
    # Expired articles -> back to draft
    expired = await db.articles.update_many(
        {
            "expiry_date": {"$lte": now, "$ne": None},
            "status": {"$ne": "draft"}
        },
        {"$set": {"status": "draft"}}
    )
    if expired.modified_count > 0:
        logger.info(f"Set {expired.modified_count} expired articles to draft")
    
    # Expired important markings
    important_expired = await db.articles.update_many(
        {
            "important_until": {"$lte": now, "$ne": None},
            "is_important": True
        },
        {"$set": {"is_important": False, "important_until": None}}
    )
    if important_expired.modified_count > 0:
        logger.info(f"Removed important marking from {important_expired.modified_count} articles")

@app.on_event("shutdown")
async def shutdown_db_client():
    client.close()

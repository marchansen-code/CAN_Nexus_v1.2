"""
Document upload and processing routes for the CANUSA Knowledge Hub API.
Supports PDF, DOC/DOCX, TXT, CSV, XLS/XLSX files.
"""
from fastapi import APIRouter, HTTPException, Depends, File, UploadFile, Form
from fastapi.responses import FileResponse, Response
from typing import Dict, List
from datetime import datetime, timezone
import asyncio
import uuid
import os
import logging
import pdfplumber
import json

from database import db
from models import User, Document
from dependencies import get_current_user

router = APIRouter(prefix="/documents", tags=["Documents"])
logger = logging.getLogger(__name__)

# Supported file extensions
SUPPORTED_EXTENSIONS = {
    '.pdf': 'application/pdf',
    '.doc': 'application/msword',
    '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    '.txt': 'text/plain',
    '.csv': 'text/csv',
    '.xls': 'application/vnd.ms-excel',
    '.xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
}

def get_file_extension(filename: str) -> str:
    """Get lowercase file extension."""
    return os.path.splitext(filename.lower())[1]


def is_supported_file(filename: str) -> bool:
    """Check if file type is supported."""
    return get_file_extension(filename) in SUPPORTED_EXTENSIONS


async def process_document_content(file_path: str, mime_type: str) -> dict:
    """Process document and return structured content. Used by Google Drive import."""
    ext_map = {
        'application/pdf': '.pdf',
        'application/msword': '.doc',
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document': '.docx',
        'text/plain': '.txt',
        'text/csv': '.csv',
        'application/vnd.ms-excel': '.xls',
        'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': '.xlsx',
    }
    
    file_type = ext_map.get(mime_type, '.txt')
    
    try:
        if file_type == '.pdf':
            result = await process_pdf(file_path)
        elif file_type in ['.doc', '.docx']:
            result = await process_word(file_path)
        elif file_type == '.txt':
            result = await process_text(file_path)
        elif file_type in ['.csv', '.xls', '.xlsx']:
            result = await process_spreadsheet(file_path, file_type)
        else:
            result = {"extracted_text": "", "html_content": "", "page_count": 0}
        
        return {
            "extracted_text": result.get("extracted_text", ""),
            "html_content": result.get("html_content", ""),
            "page_count": result.get("page_count", 0)
        }
    except Exception as e:
        logger.error(f"Content processing failed: {e}")
        return {"extracted_text": "", "html_content": "", "page_count": 0}


@router.post("/upload")
async def upload_document(
    file: UploadFile = File(...),
    target_language: str = Form("de"),
    folder_id: str = Form(None),
    force: bool = Form(False),
    user: User = Depends(get_current_user)
):
    """Upload a document for processing. Supports PDF, DOC/DOCX, TXT, CSV, XLS/XLSX."""
    ext = get_file_extension(file.filename)
    
    if not is_supported_file(file.filename):
        supported = ", ".join(SUPPORTED_EXTENSIONS.keys())
        raise HTTPException(
            status_code=400, 
            detail=f"Dateityp nicht unterstützt. Erlaubte Formate: {supported}"
        )
    
    existing_doc = await db.documents.find_one({"filename": file.filename})
    if existing_doc and not force:
        raise HTTPException(
            status_code=409, 
            detail=f"Eine Datei mit dem Namen '{file.filename}' existiert bereits."
        )
    
    if existing_doc and force:
        old_path = existing_doc.get("file_path")
        if old_path and os.path.exists(old_path):
            os.remove(old_path)
        await db.documents.delete_one({"document_id": existing_doc["document_id"]})
    
    # Verify folder exists if provided
    if folder_id:
        folder = await db.document_folders.find_one({"folder_id": folder_id})
        if not folder:
            raise HTTPException(status_code=404, detail="Ordner nicht gefunden")
    
    content = await file.read()
    doc_id = f"doc_{uuid.uuid4().hex[:12]}"
    permanent_path = f"/tmp/docs/{doc_id}{ext}"
    
    os.makedirs("/tmp/docs", exist_ok=True)
    
    with open(permanent_path, "wb") as f:
        f.write(content)
    
    doc = Document(
        document_id=doc_id,
        filename=file.filename,
        target_language=target_language,
        status="pending",
        uploaded_by=user.user_id
    )
    doc_dict = doc.model_dump()
    doc_dict["created_at"] = doc_dict["created_at"].isoformat()
    doc_dict["file_path"] = permanent_path
    doc_dict["file_size"] = len(content)
    doc_dict["folder_id"] = folder_id
    doc_dict["file_type"] = ext
    
    await db.documents.insert_one(doc_dict)
    
    asyncio.create_task(process_document(doc_id, permanent_path, target_language, ext))
    
    return {
        "document_id": doc_id,
        "filename": doc.filename,
        "folder_id": folder_id,
        "file_type": ext,
        "status": "pending",
        "message": "Dokument wird verarbeitet"
    }


async def process_document(document_id: str, file_path: str, target_language: str, file_type: str):
    """Process document based on file type."""
    try:
        await db.documents.update_one(
            {"document_id": document_id},
            {"$set": {"status": "processing"}}
        )
        
        if file_type == '.pdf':
            result = await process_pdf(file_path)
        elif file_type in ['.doc', '.docx']:
            result = await process_word(file_path)
        elif file_type == '.txt':
            result = await process_text(file_path)
        elif file_type in ['.csv', '.xls', '.xlsx']:
            result = await process_spreadsheet(file_path, file_type)
        else:
            result = {"extracted_text": "", "html_content": "", "page_count": 0}
        
        await db.documents.update_one(
            {"document_id": document_id},
            {"$set": {
                "status": "completed",
                "extracted_text": result.get("extracted_text", ""),
                "html_content": result.get("html_content", ""),
                "page_count": result.get("page_count", 0),
                "processed_at": datetime.now(timezone.utc).isoformat()
            }}
        )
        
    except Exception as e:
        logger.error(f"Document processing failed: {e}")
        await db.documents.update_one(
            {"document_id": document_id},
            {"$set": {
                "status": "failed",
                "error": str(e)
            }}
        )


async def process_pdf(file_path: str) -> dict:
    """Process PDF document: extract text and tables."""
    extracted_text = ""
    html_content = ""
    page_count = 0
    
    try:
        with pdfplumber.open(file_path) as pdf:
            page_count = len(pdf.pages)
            for page_num, page in enumerate(pdf.pages, 1):
                html_content += f"<div class='pdf-page' data-page='{page_num}'>"
                
                text = page.extract_text()
                if text:
                    extracted_text += f"--- Seite {page_num} ---\n{text}\n\n"
                    paragraphs = text.split('\n\n')
                    for para in paragraphs:
                        if para.strip():
                            if len(para.strip()) < 100 and para.strip().isupper():
                                html_content += f"<h3>{para.strip()}</h3>"
                            else:
                                html_content += f"<p>{para.strip()}</p>"
                
                tables = page.extract_tables()
                for table_idx, table in enumerate(tables):
                    if table and len(table) > 0:
                        table_html = "<table class='w-full border-collapse my-4 text-sm'>"
                        for row_idx, row in enumerate(table):
                            if row:
                                tag = 'th' if row_idx == 0 else 'td'
                                cell_class = 'border border-slate-300 p-2 bg-slate-50 font-medium' if row_idx == 0 else 'border border-slate-300 p-2'
                                table_html += "<tr>"
                                for cell in row:
                                    cell_text = str(cell) if cell else ""
                                    table_html += f"<{tag} class='{cell_class}'>{cell_text}</{tag}>"
                                table_html += "</tr>"
                        table_html += "</table>"
                        html_content += table_html
                
                html_content += "</div>"
    except Exception as e:
        logger.error(f"PDF processing error: {e}")
    
    return {"extracted_text": extracted_text, "html_content": html_content, "page_count": page_count}


async def process_word(file_path: str) -> dict:
    """Process Word document (DOC/DOCX)."""
    from docx import Document as DocxDocument
    import html as html_module
    
    extracted_text = ""
    html_content = ""
    
    try:
        doc = DocxDocument(file_path)
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                extracted_text += text + "\n\n"
                # Escape HTML entities to prevent rendering issues
                escaped_text = html_module.escape(text)
                
                # Check style for headings
                if para.style and para.style.name and para.style.name.startswith('Heading'):
                    level = para.style.name[-1] if para.style.name[-1].isdigit() else '2'
                    html_content += f"<h{level}>{escaped_text}</h{level}>"
                else:
                    html_content += f"<p>{escaped_text}</p>"
        
        # Process tables
        for table in doc.tables:
            table_html = "<table class='w-full border-collapse my-4 text-sm'>"
            for row_idx, row in enumerate(table.rows):
                tag = 'th' if row_idx == 0 else 'td'
                cell_class = 'border border-slate-300 p-2 bg-slate-50 font-medium' if row_idx == 0 else 'border border-slate-300 p-2'
                table_html += "<tr>"
                for cell in row.cells:
                    cell_text = html_module.escape(cell.text.strip())
                    table_html += f"<{tag} class='{cell_class}'>{cell_text}</{tag}>"
                table_html += "</tr>"
            table_html += "</table>"
            html_content += table_html
            extracted_text += "[Tabelle]\n\n"
        
        page_count = max(1, len(doc.paragraphs) // 30)  # Estimate pages
        
    except Exception as e:
        logger.error(f"Word processing error: {e}")
        return {"extracted_text": "", "html_content": "", "page_count": 0}
    
    return {"extracted_text": extracted_text, "html_content": html_content, "page_count": page_count}


async def process_text(file_path: str) -> dict:
    """Process plain text file."""
    extracted_text = ""
    html_content = ""
    
    try:
        with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
            content = f.read()
        
        extracted_text = content
        
        # Convert to HTML with paragraph breaks
        paragraphs = content.split('\n\n')
        for para in paragraphs:
            if para.strip():
                # Check if it looks like a heading (short, possibly uppercase)
                lines = para.strip().split('\n')
                for line in lines:
                    if line.strip():
                        if len(line.strip()) < 80 and (line.strip().isupper() or line.strip().startswith('#')):
                            clean_line = line.strip().lstrip('#').strip()
                            html_content += f"<h3>{clean_line}</h3>"
                        else:
                            html_content += f"<p>{line.strip()}</p>"
        
        page_count = max(1, len(content) // 3000)  # Estimate pages
        
    except Exception as e:
        logger.error(f"Text processing error: {e}")
        return {"extracted_text": "", "html_content": "", "page_count": 0}
    
    return {"extracted_text": extracted_text, "html_content": html_content, "page_count": page_count}


async def process_spreadsheet(file_path: str, file_type: str) -> dict:
    """Process spreadsheet (CSV, XLS, XLSX)."""
    import pandas as pd
    
    extracted_text = ""
    html_content = ""
    page_count = 0
    
    try:
        if file_type == '.csv':
            # Try different encodings
            for encoding in ['utf-8', 'latin-1', 'cp1252']:
                try:
                    df = pd.read_csv(file_path, encoding=encoding)
                    break
                except Exception:
                    continue
            else:
                df = pd.read_csv(file_path, encoding='utf-8', errors='replace')
        elif file_type == '.xls':
            df = pd.read_excel(file_path, engine='xlrd')
        else:  # xlsx
            df = pd.read_excel(file_path, engine='openpyxl')
        
        # Build HTML table
        table_html = "<table class='w-full border-collapse my-4 text-sm'>"
        
        # Header
        table_html += "<thead><tr>"
        for col in df.columns:
            table_html += f"<th class='border border-slate-300 p-2 bg-slate-100 font-medium text-left'>{col}</th>"
        table_html += "</tr></thead>"
        
        # Body
        table_html += "<tbody>"
        for idx, row in df.iterrows():
            table_html += "<tr>"
            for val in row:
                cell_val = str(val) if pd.notna(val) else ""
                table_html += f"<td class='border border-slate-300 p-2'>{cell_val}</td>"
            table_html += "</tr>"
        table_html += "</tbody></table>"
        
        html_content = table_html
        
        # Build text representation
        extracted_text = df.to_string(index=False)
        
        page_count = max(1, len(df) // 50)  # Estimate pages
        
    except Exception as e:
        logger.error(f"Spreadsheet processing error: {e}")
        return {"extracted_text": "", "html_content": "", "page_count": 0}
    
    return {"extracted_text": extracted_text, "html_content": html_content, "page_count": page_count}


@router.get("", response_model=List[Dict])
async def get_documents(user: User = Depends(get_current_user)):
    """Get all documents (excluding soft-deleted)."""
    docs = await db.documents.find(
        {"deleted_at": {"$exists": False}}, 
        {"_id": 0, "temp_path": 0}
    ).sort("created_at", -1).to_list(100)
    return docs


@router.get("/{document_id}", response_model=Dict)
async def get_document(document_id: str, user: User = Depends(get_current_user)):
    """Get document details."""
    doc = await db.documents.find_one(
        {"document_id": document_id, "deleted_at": {"$exists": False}}, 
        {"_id": 0, "temp_path": 0}
    )
    if not doc:
        raise HTTPException(status_code=404, detail="Dokument nicht gefunden")
    return doc


@router.delete("/{document_id}")
async def delete_document(document_id: str, user: User = Depends(get_current_user)):
    """Soft delete a document - moves to trash for 30 days (admin and editor only)."""
    if user.role not in ["admin", "editor"]:
        raise HTTPException(status_code=403, detail="Nur Administratoren und Editoren können Dokumente löschen")
    
    doc = await db.documents.find_one({"document_id": document_id})
    if not doc:
        raise HTTPException(status_code=404, detail="Dokument nicht gefunden")
    
    await db.documents.update_one(
        {"document_id": document_id},
        {"$set": {
            "deleted_at": datetime.now(timezone.utc),
            "deleted_by": user.user_id
        }}
    )
    return {"message": "Dokument in Papierkorb verschoben"}


@router.get("/{document_id}/file")
async def get_document_file(document_id: str, inline: bool = False, user: User = Depends(get_current_user)):
    """Get the actual file for viewing or downloading."""
    doc = await db.documents.find_one(
        {"document_id": document_id, "deleted_at": {"$exists": False}},
        {"_id": 0}
    )
    if not doc:
        raise HTTPException(status_code=404, detail="Dokument nicht gefunden")
    
    file_path = doc.get("file_path")
    if not file_path or not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="Datei nicht gefunden")
    
    # Get file type and appropriate media type
    file_type = doc.get("file_type", ".pdf")
    media_type = SUPPORTED_EXTENSIONS.get(file_type, "application/octet-stream")
    
    # For inline viewing (preview), don't set filename to avoid download
    if inline:
        return FileResponse(
            path=file_path,
            media_type=media_type
        )
    else:
        return FileResponse(
            path=file_path,
            media_type=media_type,
            filename=doc.get("filename", "document")
        )


@router.get("/{document_id}/preview")
async def get_document_preview(document_id: str, user: User = Depends(get_current_user)):
    """Get the document file for inline preview (no download prompt)."""
    doc = await db.documents.find_one(
        {"document_id": document_id, "deleted_at": {"$exists": False}},
        {"_id": 0}
    )
    if not doc:
        raise HTTPException(status_code=404, detail="Dokument nicht gefunden")
    
    file_path = doc.get("file_path")
    if not file_path or not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="Datei nicht gefunden")
    
    # Get file type and appropriate media type
    file_type = doc.get("file_type", ".pdf")
    media_type = SUPPORTED_EXTENSIONS.get(file_type, "application/octet-stream")
    
    # Return file without Content-Disposition: attachment header
    with open(file_path, "rb") as f:
        content = f.read()
    
    return Response(
        content=content,
        media_type=media_type,
        headers={
            "Content-Disposition": f"inline; filename=\"{doc.get('filename', 'document')}\""
        }
    )


@router.get("/{document_id}/content")
async def get_document_content(document_id: str, user: User = Depends(get_current_user)):
    """Get the processed HTML content of a document for embedding in articles."""
    doc = await db.documents.find_one(
        {"document_id": document_id, "deleted_at": {"$exists": False}},
        {"_id": 0}
    )
    if not doc:
        raise HTTPException(status_code=404, detail="Dokument nicht gefunden")
    
    if doc.get("status") != "completed":
        raise HTTPException(status_code=400, detail="Dokument muss zuerst verarbeitet werden")
    
    return {
        "document_id": document_id,
        "filename": doc.get("filename"),
        "file_type": doc.get("file_type"),
        "html_content": doc.get("html_content") or (doc.get("structured_content") or {}).get("html_content", ""),
        "extracted_text": doc.get("extracted_text", "")
    }


@router.put("/{document_id}/move")
async def move_document(document_id: str, folder_id: str = None, user: User = Depends(get_current_user)):
    """Move a document to a different folder."""
    if user.role not in ["admin", "editor"]:
        raise HTTPException(status_code=403, detail="Nur Editoren und Administratoren können Dokumente verschieben")
    
    doc = await db.documents.find_one({"document_id": document_id})
    if not doc:
        raise HTTPException(status_code=404, detail="Dokument nicht gefunden")
    
    # Verify folder exists if provided
    if folder_id:
        folder = await db.document_folders.find_one({"folder_id": folder_id})
        if not folder:
            raise HTTPException(status_code=404, detail="Ordner nicht gefunden")
    
    await db.documents.update_one(
        {"document_id": document_id},
        {"$set": {"folder_id": folder_id}}
    )
    
    return {"message": "Dokument verschoben", "folder_id": folder_id}



@router.get("/{document_id}/convert-to-html")
async def convert_pdf_to_html(document_id: str, user: User = Depends(get_current_user)):
    """Convert PDF document to editable HTML content for the TipTap editor."""
    doc = await db.documents.find_one(
        {"document_id": document_id, "deleted_at": {"$exists": False}}, 
        {"_id": 0}
    )
    if not doc:
        raise HTTPException(status_code=404, detail="Dokument nicht gefunden")
    
    if doc.get("status") != "completed":
        raise HTTPException(status_code=400, detail="Dokument muss zuerst verarbeitet werden")
    
    file_path = doc.get("file_path")
    if not file_path or not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="PDF-Datei nicht gefunden")
    
    try:
        import pymupdf4llm
        import markdown
        
        # Extract markdown from PDF with table support
        md_text = pymupdf4llm.to_markdown(file_path)
        
        # Convert markdown to HTML
        html_content = markdown.markdown(
            md_text,
            extensions=['tables', 'fenced_code', 'nl2br']
        )
        
        # Post-process HTML for TipTap compatibility
        html_content = post_process_html_for_tiptap(html_content)
        
        return {
            "document_id": document_id,
            "filename": doc.get("filename"),
            "html_content": html_content,
            "success": True
        }
        
    except Exception as e:
        logger.error(f"PDF to HTML conversion failed: {e}")
        raise HTTPException(status_code=500, detail=f"Konvertierung fehlgeschlagen: {str(e)}")


def post_process_html_for_tiptap(html: str) -> str:
    """Post-process HTML to be compatible with TipTap editor."""
    import re
    
    # Wrap tables in a div for better styling
    html = re.sub(
        r'<table>',
        '<table class="w-full border-collapse my-4">',
        html
    )
    
    # Add styling to table cells
    html = re.sub(
        r'<th>',
        '<th class="border border-slate-400 bg-slate-100 p-2 font-semibold text-left">',
        html
    )
    html = re.sub(
        r'<td>',
        '<td class="border border-slate-300 p-2">',
        html
    )
    
    # Convert h1-h6 headings to proper format
    for i in range(1, 7):
        html = re.sub(
            rf'<h{i}>',
            f'<h{i} style="margin-top: 1rem; margin-bottom: 0.5rem;">',
            html
        )
    
    # Ensure paragraphs have proper spacing
    html = re.sub(
        r'<p>',
        '<p style="margin-bottom: 0.75rem;">',
        html
    )
    
    return html

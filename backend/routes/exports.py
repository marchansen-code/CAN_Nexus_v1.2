"""
Article export routes (PDF, DOCX) and miscellaneous endpoints.
"""
from fastapi import APIRouter, HTTPException, Depends
from fastapi.responses import StreamingResponse, Response
from datetime import datetime, timezone, timedelta
import re
import io
import os

from database import db
from models import User, UserGroupUpdate
from dependencies import get_current_user

router = APIRouter(tags=["Exports"])

# In-memory storage for active editors
active_editors = {}


def strip_html(html_content: str) -> str:
    """Remove HTML tags and convert to plain text."""
    text = re.sub(r'<script[^>]*>.*?</script>', '', html_content, flags=re.DOTALL)
    text = re.sub(r'<style[^>]*>.*?</style>', '', text, flags=re.DOTALL)
    text = re.sub(r'<br\s*/?>', '\n', text)
    text = re.sub(r'</p>', '\n\n', text)
    text = re.sub(r'</h[1-6]>', '\n\n', text)
    text = re.sub(r'<li>', '• ', text)
    text = re.sub(r'</li>', '\n', text)
    text = re.sub(r'<[^>]+>', '', text)
    text = re.sub(r'\n\s*\n', '\n\n', text)
    text = text.strip()
    return text


@router.get("/articles/{article_id}/export/pdf")
async def export_article_pdf(article_id: str, user: User = Depends(get_current_user)):
    """Export article as PDF."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER
    
    article = await db.articles.find_one({"article_id": article_id}, {"_id": 0})
    if not article:
        raise HTTPException(status_code=404, detail="Artikel nicht gefunden")
    
    category_name = "Keine Kategorie"
    if article.get("category_id"):
        cat = await db.categories.find_one({"category_id": article["category_id"]}, {"_id": 0, "name": 1})
        if cat:
            category_name = cat["name"]
    
    author_name = "Unbekannt"
    if article.get("created_by"):
        author = await db.users.find_one({"user_id": article["created_by"]}, {"_id": 0, "name": 1})
        if author:
            author_name = author["name"]
    
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=2*cm,
        leftMargin=2*cm,
        topMargin=2*cm,
        bottomMargin=2*cm
    )
    
    styles = getSampleStyleSheet()
    
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=18,
        spaceAfter=12,
        textColor='#1e293b'
    )
    
    meta_style = ParagraphStyle(
        'Meta',
        parent=styles['Normal'],
        fontSize=10,
        textColor='#64748b',
        spaceAfter=20
    )
    
    body_style = ParagraphStyle(
        'CustomBody',
        parent=styles['Normal'],
        fontSize=11,
        leading=16,
        alignment=TA_JUSTIFY,
        spaceAfter=12
    )
    
    story = []
    
    story.append(Paragraph(article["title"], title_style))
    
    created_at = article.get("created_at", "")
    if isinstance(created_at, str) and created_at:
        try:
            dt = datetime.fromisoformat(created_at.replace('Z', '+00:00'))
            created_at = dt.strftime("%d.%m.%Y")
        except Exception:
            pass
    
    meta_text = f"Kategorie: {category_name} | Autor: {author_name} | Erstellt: {created_at}"
    story.append(Paragraph(meta_text, meta_style))
    
    if article.get("summary"):
        story.append(Paragraph(f"<b>Zusammenfassung:</b> {article['summary']}", body_style))
        story.append(Spacer(1, 12))
    
    content = strip_html(article.get("content", ""))
    paragraphs = content.split('\n\n')
    for para in paragraphs:
        if para.strip():
            para = para.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            story.append(Paragraph(para, body_style))
    
    story.append(Spacer(1, 30))
    footer_style = ParagraphStyle('Footer', parent=styles['Normal'], fontSize=8, textColor='#94a3b8', alignment=TA_CENTER)
    story.append(Paragraph(f"Exportiert aus CANUSA Nexus am {datetime.now().strftime('%d.%m.%Y %H:%M')}", footer_style))
    
    doc.build(story)
    buffer.seek(0)
    
    safe_title = "".join(c for c in article["title"] if c.isalnum() or c in (' ', '-', '_')).strip()[:50]
    filename = f"{safe_title}.pdf"
    
    return StreamingResponse(
        buffer,
        media_type="application/pdf",
        headers={
            "Content-Disposition": f"attachment; filename=\"{filename}\""
        }
    )


@router.get("/articles/{article_id}/export/docx")
async def export_article_docx(article_id: str, user: User = Depends(get_current_user)):
    """Export article as Word document."""
    from docx import Document as DocxDocument
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    article = await db.articles.find_one({"article_id": article_id}, {"_id": 0})
    if not article:
        raise HTTPException(status_code=404, detail="Artikel nicht gefunden")
    
    category_name = "Keine Kategorie"
    if article.get("category_id"):
        cat = await db.categories.find_one({"category_id": article["category_id"]}, {"_id": 0, "name": 1})
        if cat:
            category_name = cat["name"]
    
    author_name = "Unbekannt"
    if article.get("created_by"):
        author = await db.users.find_one({"user_id": article["created_by"]}, {"_id": 0, "name": 1})
        if author:
            author_name = author["name"]
    
    doc = DocxDocument()
    
    title = doc.add_heading(article["title"], level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    created_at = article.get("created_at", "")
    if isinstance(created_at, str) and created_at:
        try:
            dt = datetime.fromisoformat(created_at.replace('Z', '+00:00'))
            created_at = dt.strftime("%d.%m.%Y")
        except Exception:
            pass
    
    meta = doc.add_paragraph()
    meta_run = meta.add_run(f"Kategorie: {category_name} | Autor: {author_name} | Erstellt: {created_at}")
    meta_run.font.size = Pt(10)
    meta_run.font.color.rgb = RGBColor(100, 116, 139)
    
    doc.add_paragraph()
    
    if article.get("summary"):
        summary_para = doc.add_paragraph()
        summary_run = summary_para.add_run("Zusammenfassung: ")
        summary_run.bold = True
        summary_para.add_run(article["summary"])
        doc.add_paragraph()
    
    content = strip_html(article.get("content", ""))
    paragraphs = content.split('\n\n')
    for para in paragraphs:
        if para.strip():
            p = doc.add_paragraph(para.strip())
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer_run = footer.add_run(f"Exportiert aus CANUSA Nexus am {datetime.now().strftime('%d.%m.%Y %H:%M')}")
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(148, 163, 184)
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    safe_title = "".join(c for c in article["title"] if c.isalnum() or c in (' ', '-', '_')).strip()[:50]
    filename = f"{safe_title}.docx"
    
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f"attachment; filename=\"{filename}\""
        }
    )


# ==================== FAVORITES & VIEWED ====================

@router.post("/articles/{article_id}/favorite")
async def toggle_favorite(article_id: str, user: User = Depends(get_current_user)):
    """Toggle favorite status for an article."""
    article = await db.articles.find_one({"article_id": article_id}, {"_id": 0})
    if not article:
        raise HTTPException(status_code=404, detail="Artikel nicht gefunden")
    
    favorited_by = article.get("favorited_by", [])
    
    if user.user_id in favorited_by:
        await db.articles.update_one(
            {"article_id": article_id},
            {"$pull": {"favorited_by": user.user_id}}
        )
        return {"favorited": False, "message": "Aus Favoriten entfernt"}
    else:
        await db.articles.update_one(
            {"article_id": article_id},
            {"$addToSet": {"favorited_by": user.user_id}}
        )
        return {"favorited": True, "message": "Zu Favoriten hinzugefügt"}


@router.post("/articles/{article_id}/viewed")
async def mark_as_viewed(article_id: str, user: User = Depends(get_current_user)):
    """Mark article as viewed and update recently viewed list."""
    await db.users.update_one(
        {"user_id": user.user_id},
        {"$pull": {"recently_viewed": article_id}}
    )
    await db.users.update_one(
        {"user_id": user.user_id},
        {"$push": {"recently_viewed": {"$each": [article_id], "$position": 0, "$slice": 20}}}
    )
    await db.articles.update_one(
        {"article_id": article_id},
        {"$inc": {"view_count": 1}}
    )
    return {"message": "Als angesehen markiert"}


# ==================== PRESENCE / ACTIVE EDITORS ====================

@router.post("/articles/{article_id}/presence")
async def update_presence(article_id: str, user: User = Depends(get_current_user)):
    """Update editor presence for an article."""
    global active_editors
    
    if article_id not in active_editors:
        active_editors[article_id] = {}
    
    active_editors[article_id][user.user_id] = {
        "name": user.name,
        "timestamp": datetime.now(timezone.utc).isoformat()
    }
    
    cutoff = datetime.now(timezone.utc) - timedelta(seconds=30)
    for uid in list(active_editors[article_id].keys()):
        ts = datetime.fromisoformat(active_editors[article_id][uid]["timestamp"])
        if ts < cutoff:
            del active_editors[article_id][uid]
    
    others = {uid: info for uid, info in active_editors[article_id].items() if uid != user.user_id}
    
    return {"active_editors": list(others.values())}


@router.delete("/articles/{article_id}/presence")
async def remove_presence(article_id: str, user: User = Depends(get_current_user)):
    """Remove editor presence when leaving article."""
    global active_editors
    
    if article_id in active_editors and user.user_id in active_editors[article_id]:
        del active_editors[article_id][user.user_id]
    
    return {"message": "Präsenz entfernt"}


# ==================== USER GROUPS ====================

@router.put("/users/{user_id}/groups")
async def update_user_groups(user_id: str, data: UserGroupUpdate, current_user: User = Depends(get_current_user)):
    """Update user's group memberships (admin only)."""
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Nur Administratoren können Gruppenzugehörigkeiten ändern")
    
    for gid in data.group_ids:
        exists = await db.groups.find_one({"group_id": gid})
        if not exists:
            raise HTTPException(status_code=400, detail=f"Gruppe {gid} nicht gefunden")
    
    result = await db.users.update_one(
        {"user_id": user_id},
        {"$set": {"group_ids": data.group_ids}}
    )
    
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Benutzer nicht gefunden")
    
    return {"message": "Gruppenzugehörigkeiten aktualisiert", "group_ids": data.group_ids}


# ==================== PDF EMBED ====================

@router.get("/documents/{document_id}/pdf-embed")
async def get_document_pdf_embed(document_id: str, token: str = None):
    """Get PDF file for iframe embedding."""
    doc = await db.documents.find_one({"document_id": document_id})
    if not doc:
        raise HTTPException(status_code=404, detail="Dokument nicht gefunden")
    
    file_path = doc.get("file_path")
    if not file_path or not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="PDF-Datei nicht gefunden")
    
    with open(file_path, "rb") as f:
        content = f.read()
    
    return Response(
        content=content,
        media_type="application/pdf",
        headers={
            "Content-Disposition": "inline",
            "Cache-Control": "public, max-age=3600",
            "X-Frame-Options": "SAMEORIGIN"
        }
    )
